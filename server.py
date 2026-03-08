from flask import Flask, request
from flask_cors import CORS
from docx import Document
from deep_translator import GoogleTranslator
from textblob import TextBlob
import os
import threading

app = Flask(__name__)
CORS(app)

@app.route("/")
def home():
    return "Vocabulary API is running"

file_name = "wordsheet.docx"
lock = threading.Lock()

# create file if not exists
if not os.path.exists(file_name):
    doc = Document()
    doc.add_heading("Vocabulary Notebook", level=1)
    doc.save(file_name)

def word_exists(word):
    doc = Document(file_name)

    for para in doc.paragraphs:
        if para.text.lower().startswith(word + " :"):
            return True

    return False


@app.route("/save", methods=["POST"])
def save_word():

    data = request.json

    if not data or "word" not in data:
        return {"status": "no_word"}

    word = data["word"].lower().strip()
    word = word.split()[0]

    corrected_word = str(TextBlob(word).correct())

    with lock:

        if word_exists(corrected_word):
            return {"status": "duplicate"}

        translation = GoogleTranslator(source='auto', target='hi').translate(corrected_word)

        doc = Document(file_name)
        doc.add_paragraph(f"{corrected_word} : {translation}")

        try:
            doc.save(file_name)
        except PermissionError:
            return {"status": "file_open"}

        print(corrected_word, "->", translation)

    return {"status": "saved"}


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
